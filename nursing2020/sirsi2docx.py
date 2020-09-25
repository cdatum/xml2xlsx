"""

Christopher Galluzzo
May 1, 2019
Python 1 - Spring 2019
Final - SirsiDynix XML to MS Word
"""

from bs4 import BeautifulSoup
from docx import Document      # in Anaconda Prompt, ran pip install python-docx - Will allow output to MS Word file
from docx.shared import Inches # for indenting paragraphs 0.5 inches
from datetime import datetime  # for date/timestamp
import os # for os.listdir()

'''
This program is designed to convert XML reports from the SirsiDynix library system
into Microsoft Word files. While SirsiDynix reports can include any number of bibliographic 
and item record fields, this program focuses on this set of data fields:
	

<catalog>
<catalogKey>187868</catalogKey>
<yearOfPublication>2012</yearOfPublication>
<marc>
<marcEntry tag="020" label="ISBN" ind="  ">9781608317905 (pbk. : alkaline paper)</marcEntry>
<marcEntry tag="020" label="ISBN" ind="  ">1608317900 (pbk. : alkaline paper)</marcEntry>
<marcEntry tag="050" label="LC Call Number" ind="00">RS57 .C73 2012</marcEntry>
<marcEntry tag="100" label="Personal Author" ind="1 ">Craig, Gloria P., 1949-</marcEntry>
<marcEntry tag="245" label="Title" ind="10">Clinical calculations made easy : solving problems using dimensional analysis / Gloria P. Craig, RN, MSN, EdD ; Associate Professor, South Dakota State University, College of Nursing, Brookings, South Dakota.</marcEntry>
<marcEntry tag="260" label="Publication info:" ind="  ">Philadelphia : Wolters Kluwer Health/Lippincott Williams &amp; Wilkins Health, [2012], ©2012.</marcEntry>
<marcEntry tag="300" label="Physical description" ind="  ">xii, 292 pages : illustrations ; 28 cm</marcEntry>
<marcEntry tag="596" label="Held by" ind="  ">NORTH </marcEntry>
</marc>
<call>
<callNumber>RS 57 .C73 2012</callNumber>
<library>NORTH</library>
<numberOfCallHolds>0</numberOfCallHolds>
<numberOfCopies>1</numberOfCopies>
<copiesOnReserve>1</copiesOnReserve>
<item>
<numberOfCharges>0</numberOfCharges>
<numberOfCharges>0</numberOfCharges>
<numberOfBills>0</numberOfBills>
<numberOfCopyHolds>0</numberOfCopyHolds>
<totalCharges>19</totalCharges>
<inhouseCharges>0</inhouseCharges>
<totalCheckouts>17</totalCheckouts>
<totalRenewals>2</totalRenewals>
<intervalCheckouts>17</intervalCheckouts>
<intervalRenewals>2</intervalRenewals>
<intervalStartDate></intervalStartDate>
<recirculate>YES</recirculate>
<dateLastUsed>2015-06-30</dateLastUsed>
<isReserveItem>true</isReserveItem>
<copyNumber>1</copyNumber>
<itemID>3127800000938717</itemID>
<library>NORTH</library>
<libraryDescription>North Campus Library</libraryDescription>
<location>RESERVES</location>
</item>
</call>
</catalog>

HOW THIS PROGRAM WORKS
This program will batch process files within a given directory. So, if the user
has multiple reports to convert, once the user enters the directory location,
the program will find all of the XML files and then produce an MS Word file
for each.

Each MS Word file will be given a name that corresponds with the source  
XML file, but with a timestamp appended to the end. This will keep filenames unique
and prevent files from being overwritten.

USAGE
Users will be prompted to enter a path for a directory that has XML files to be
converted. If the user doesn't enter anything, the directory will be the same
as where the program is located. 


'''

# Holds the directory where xml file(s) will be sourced and docx file(s) will be saved
directory_path = ''
# Flag for when no XML files are found
no_xml = True

def get_catalog_details(tag):
    print("name:", tag.name)
    print("attrs:", tag.attrs)

# Get a list of .xml files in the directory
def get_filelist():
    xml_filenames = []
    # https://docs.python.org/3/library/os.html#os.listdir
    # A GUI might be nice for finding the right directory
    dir = input("Enter a directory that contains xml files to be converted OR hit enter to continue with current dir. \n")
    
    global directory_path
    global no_xml
    # Use the current directory if the user inputs nothing
    if dir == '':
        directory_path = os.getcwd()
        files_list = os.listdir()
        print("Searching", os.getcwd()) # https://docs.python.org/3/library/os.html#os.getcwd
        # Otherwise use the user-specified directory
    else:
        directory_path = dir
        files_list = os.listdir(dir)
        print("Searching", dir)
    
    for file in files_list:
        if file.endswith('.xml'):
            xml_filenames.append(file)
            
    
    # Let user know how many files were found
    if len(xml_filenames) == 0:
        print("\nSorry, no files found in that directory")
        
    elif len(files_list) == 1:
        print("\nFound", len(xml_filenames), "file to convert.")
        no_xml = False
    else:
        print("\nFound", len(xml_filenames), "files to convert. ")
        no_xml = False
    

    # Iterate through each file, process it, and generate a Word doc
    for xml_file in xml_filenames:    
        convert_xml_to_word(xml_file)    
    

# This function extracts the bulk of the data from the xml file.
# Accepts an xml file and produces a Word file. Only prints select
# fields; not an entire bib or item record
def convert_xml_to_word(xml_file):
    # Initialize a new Word doc
    doc = Document()
    # Grab the source filename minus '.xml' This will be used in the title of the .docx
    docx_title = xml_file.replace('.xml', '') + '_'
    docx_heading = "Report for " + xml_file
    doc.add_heading(docx_heading, level=1)
    date_paragraph = doc.add_paragraph('')

    #doc.add_heading(str(datestamp), level=3)
    print("\nConverting", xml_file)
    
    with open(xml_file, encoding = 'utf-8') as booklist:
        global directory_path
        # a list to hold circ stats
        stats_list = []
        
        soup  = BeautifulSoup(booklist, 'lxml-xml')
        marc = soup.find_all('catalog')
        tag = soup.marcEntry
        # should probably handle an xml file that isn't from Sirisi or one that isn't formatted properly
        count = 1
        for item in marc:
            
            # find_all() returns a _list_ of tags and strings–a ResultSet object. 
            # You need to iterate over the list and look at the .foo of each one.
            # https://www.crummy.com/software/BeautifulSoup/bs4/doc/
            if item.yearOfPublication != None:
                year = item.yearOfPublication.get_text()
            else: 
                year = "Check year" # A few titles like yearbooks don't have a year!
            barcode = item.itemID.get_text()
            total_charges = item.totalCharges.get_text()
            date_last_use = item.dateLastUsed.get_text()
            call_no = item.callNumber.get_text()

            # This find_all returns a list of all the marcEntry elements
            tag = item.find_all('marcEntry')

            isbn = ''
    
            for element in tag: 
                # Title
                # if 245 in element[tag] 
                if element['tag'] == '245':
                    title = element.get_text()
                    stats_title = element.get_text()
                    title = str(count) + ". " + element.get_text() + " (" + year + ")"
                    
                # Description
                if element['tag'] == '260':
                    desc = element.get_text()
                    
                # Physical desc
                if element['tag'] == '300':                    
                    desc += " " + element.get_text()
                    
                # Campus Location    
                if element['tag'] == '596':
                    campus = element.get_text()
                
                # ISBNS
                if element['tag'] == '020':
                    num  = element.get_text() + ' | '
                    isbn += num
            
            # Add bibliographic details to Word file
            doc.add_heading(title, level=2)
            doc.add_paragraph(barcode + "\t" + campus + " " + call_no + "\n" + isbn).paragraph_format.left_indent = Inches(0.25)
            doc.add_paragraph(desc).paragraph_format.left_indent = Inches(0.25)
            doc.add_paragraph("Total charges (checkout + renewals): " + total_charges + "\tDate of last use: " + date_last_use ).paragraph_format.left_indent = Inches(0.25)
            
            # Collect some basic data for item stats 
            item_stats = []
            elements = (barcode, stats_title, total_charges)
            item_stats.extend(elements)
            # Save item stats in a list for the top ten list
            stats_list.append(item_stats)
            
            count += 1
        
        # Display some stats for the user to see while program is running
        print(str(count), "Total items\t", get_circ_stats(stats_list), "Total checkouts")
        
        # Add date and brief summary of items        
        date_paragraph.insert_paragraph_before("Report created on " + soup.dateCreated.get_text())
        
        number_of_items = [str(count), "total items in this report.\t" , str(get_circ_stats(stats_list)), " total checkouts."]
        subtitle_data = ' '        
        date_paragraph.insert_paragraph_before(subtitle_data.join(number_of_items))
        
        # Add table at the end of the report that displays top 10 items that circulated
        get_top(stats_list, doc)
        
        # Add the docx_title to the file extenstion
        doc_name = docx_title + get_file_extension()
        doc.save(directory_path + '/' + doc_name)
        
        print("Done!")

# Returns a total # of checkouts for all titles within an XML file
# Accepts a list of items as a parameter; returns an int of total charges (checkouts + renewals)
def get_circ_stats(item_list):
    sort_list = []
    sort_list = sorted(item_list, key=lambda record: int(record[2]), reverse=True)
    # from https://docs.python.org/3/howto/sorting.html#sortinghowto

    charges = 0
    for item in sort_list:
        charges += int(item[2])

    return charges

# Returns a table of the top 10 circulating items from the provided list
# The output is formatted as a python_docx table. Accepts 
def get_top(item_list, doc):
    sort_list = []
    sort_list = sorted(item_list, key=lambda record: int(record[2]), reverse=True)
    # from https://docs.python.org/3/howto/sorting.html#sortinghowto
    # print("Top 10 Title", sort_list[:10])
    doc.add_page_break()
    doc.add_heading("Top 20 Circulating Items in this Report", level=1)
    table = doc.add_table(rows=1, cols=3)
    # Set the col width
    set_column_width(table.columns[0], Inches(0.75))
    set_column_width(table.columns[1], Inches(4.0))
    set_column_width(table.columns[2], Inches(1.4))    

    header_row = table.rows[0].cells
    header_row[0].text = "Total Charges"
    header_row[1].text = "Title"
    header_row[2].text = "Barcode"
    
    for barcode, title, total_charges in sort_list[:20]:
        row = table.add_row().cells
        row[0].text = total_charges
        row[1].text = title
        row[2].text = barcode

    return table

    
    
# Get a timestap and use it to create a unique filename    
def get_file_extension():
    date_extension = datetime.today()
    date_extension = str(date_extension).replace('-','_')
    date_extension = date_extension.replace(':', '')
    date_extension = date_extension.replace(' ', '-')
    date_extension = date_extension.replace('.', '')
    date_extension = date_extension + '.docx'
    return date_extension


# From https://github.com/python-openxml/python-docx/issues/78#issuecomment-78421499
def set_column_width(column, width):
    column.width = width
    for cell in column.cells:
        cell.width = width

def say_goodbye():
    if no_xml == False:
        print("\nAll done! \n\nFind the Word file(s)in the following directory:\n", directory_path )
    else:
        print("\nLet's try again" )
        get_filelist()
        #TODO add a while loop for better UX 
                

#Let user enter the name of a directory
file_list = get_filelist()

say_goodbye()